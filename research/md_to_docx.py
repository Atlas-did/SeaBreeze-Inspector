#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
md_to_docx.py —— Markdown 项目报告 → 中文排版规范的 Word 文档

特性：
  * 中文排版规范：正文宋体小四（12pt）、标题黑体、1.5 倍行距、首行缩进 2 字符
  * 标题分级（## → 一级标题，### → 二级 …），自动编号可选
  * Markdown 表格 → Word 表格（表头加粗 + 浅蓝底纹）
  * 代码块 → 等宽字体 + 浅灰底纹
  * 列表 / 引用 / 加粗 / 行内代码 / 链接 完整转换
  * 自动生成目录域（在 Word 中按 Ctrl+A 后 F9 更新）与页脚页码
  * A4 页面、标准页边距

用法：
    python scripts/md_to_docx.py 输入.md [输出.docx]

示例：
    python scripts/md_to_docx.py 项目总结报告.md 项目总结报告.docx
"""
import argparse
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

# ── 排版常量（可按需修改）───────────────────────────────────────────────
BODY_CN = "宋体"          # 正文中文字体
HEAD_CN = "黑体"          # 标题中文字体
LATIN = "Times New Roman" # 西文字体
CODE_FONT = "Consolas"    # 代码字体
BODY_SIZE = 12            # 正文 小四
H1_SIZE, H2_SIZE, H3_SIZE, H4_SIZE = 16, 14, 12, 12
CODE_SIZE = 9
TABLE_SIZE = 10.5         # 表格 五号
HEADER_FILL = "D9E2F3"    # 表头底纹（浅蓝）
CODE_FILL = "F2F2F2"      # 代码底纹（浅灰）
HEAD_COLOR = RGBColor(0x1F, 0x38, 0x64)


def set_cn_font(run, cn=BODY_CN, latin=LATIN, size=BODY_SIZE, bold=None, color=None):
    """同时设置西文字体与中文（eastAsia）字体。"""
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), cn)
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_paragraph(p, fill):
    """给段落加底纹。"""
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def shade_cell(cell, fill):
    """给单元格加底纹。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_first_line_chars(p, chars=200):
    """首行缩进（按字符数，200 = 2 字符）。"""
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:firstLineChars"), str(chars))
    ind.set(qn("w:firstLine"), "480")


def add_hyperlink(p, url, text):
    """真正的可点击超链接。"""
    part = p.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                          is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    p._p.append(hyperlink)


def add_page_number_footer(doc):
    """页脚居中页码（PAGE 域）。"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    set_cn_font(run, size=9)


def add_toc(doc):
    """插入目录域（Word 中按 F9 更新）。"""
    p = doc.add_paragraph()
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "【目录】请在 Word 中选中此处后按 F9 更新域以生成目录"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(sep); run._r.append(t); run._r.append(fld2)
    return p


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
URL_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_inline(p, text, size=BODY_SIZE, cn=BODY_CN, bold=False, color=None):
    """解析行内 **加粗** / `代码` / [链接](url)。"""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            run = p.add_run(text[pos:m.start()])
            set_cn_font(run, cn=cn, size=size, bold=bold, color=color)
        tok = m.group(0)
        if tok.startswith("**"):
            run = p.add_run(tok[2:-2])
            set_cn_font(run, cn=cn, size=size, bold=True, color=color)
        elif tok.startswith("`"):
            run = p.add_run(tok[1:-1])
            set_cn_font(run, cn=CODE_FONT, latin=CODE_FONT, size=size - 1, color=RGBColor(0x9C, 0x27, 0x2A))
        else:  # link
            mm = URL_RE.match(tok)
            label, url = mm.group(1), mm.group(2)
            if url.startswith("http"):
                add_hyperlink(p, url, label)
            else:
                run = p.add_run(f"{label}（{url}）")
                set_cn_font(run, cn=cn, size=size, color=RGBColor(0x59, 0x59, 0x59))
        pos = m.end()
    if pos < len(text):
        run = p.add_run(text[pos:])
        set_cn_font(run, cn=cn, size=size, bold=bold, color=color)


def parse_table_row(line):
    """把 '| a | b |' 拆成单元格列表（去掉首尾空列）。"""
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_sep_row(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def setup_styles(doc):
    """全局样式：正文、标题、列表。"""
    normal = doc.styles["Normal"]
    normal.font.name = LATIN
    normal.font.size = Pt(BODY_SIZE)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)

    for name, size in (("Heading 1", H1_SIZE), ("Heading 2", H2_SIZE),
                       ("Heading 3", H3_SIZE), ("Heading 4", H4_SIZE)):
        st = doc.styles[name]
        st.font.name = LATIN
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = HEAD_COLOR
        st._element.rPr.rFonts.set(qn("w:eastAsia"), HEAD_CN)
        st.paragraph_format.space_before = Pt(14 if size >= 14 else 10)
        st.paragraph_format.space_after = Pt(8)
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        # 取消标题默认的“与下段同页”之外的问题保留默认


def convert(md_path, docx_path):
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    # A4 页面
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin = sec.bottom_margin = Mm(25.4)
    sec.left_margin = sec.right_margin = Mm(31.7)
    setup_styles(doc)
    add_page_number_footer(doc)

    i, n = 0, len(lines)
    code_buf = None  # 代码块缓冲
    title_done = False
    while i < n:
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        # 代码块
        if stripped.startswith("```"):
            if code_buf is None:
                code_buf = []
            else:
                p = doc.add_paragraph()
                shade_paragraph(p, CODE_FILL)
                p.paragraph_format.space_after = Pt(6)
                for j, cl in enumerate(code_buf):
                    run = p.add_run(cl if j == 0 else "\n" + cl)
                    set_cn_font(run, cn=CODE_FONT, latin=CODE_FONT, size=CODE_SIZE)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_buf = None
            i += 1
            continue
        if code_buf is not None:
            code_buf.append(raw)
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1 and not title_done:
                # 封面标题：居中大字
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(120)
                p.paragraph_format.space_after = Pt(18)
                add_inline(p, text, size=22, cn=HEAD_CN, bold=True, color=HEAD_COLOR)
                # 紧随其后的目录域（Word 中 Ctrl+A → F9 更新）
                toc_head = doc.add_paragraph(style="Heading 1")
                add_inline(toc_head, "目录", size=H1_SIZE, cn=HEAD_CN, bold=True, color=HEAD_COLOR)
                add_toc(doc)
                doc.add_page_break()
                title_done = True
            else:
                style = f"Heading {min(level, 4)}"
                p = doc.add_paragraph(style=style)
                add_inline(p, text, size={1: H1_SIZE, 2: H2_SIZE, 3: H3_SIZE, 4: H4_SIZE}[min(level, 4)],
                           cn=HEAD_CN, bold=True, color=HEAD_COLOR)
            i += 1
            continue

        # 分隔线：插一个分页符（章节分页）后的间隔段落
        if re.match(r"^---+$", line):
            if title_done:
                doc.add_page_break()
            i += 1
            continue

        # 表格
        if stripped.startswith("|") and i + 1 < n and is_sep_row(lines[i + 1].strip()):
            header = parse_table_row(line)
            j = i + 2
            rows = []
            while j < n and lines[j].strip().startswith("|"):
                rows.append(parse_table_row(lines[j]))
                j += 1
            ncols = max([len(header)] + [len(r) for r in rows])
            table = doc.add_table(rows=1 + len(rows), cols=ncols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # 表头
            for c, htxt in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                add_inline(p, htxt, size=TABLE_SIZE, bold=True)
                for r in p.runs:
                    set_cn_font(r, size=TABLE_SIZE, bold=True)
                shade_cell(cell, HEADER_FILL)
            # 数据行
            for ri, rrow in enumerate(rows, start=1):
                for c in range(ncols):
                    cell = table.rows[ri].cells[c]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_inline(p, rrow[c] if c < len(rrow) else "", size=TABLE_SIZE)
                    for r in p.runs:
                        set_cn_font(r, size=TABLE_SIZE)
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
            i = j
            continue

        # 列表
        m = re.match(r"^\s*([-*+]|\d+[.、)])\s+(.*)$", line)
        if m:
            marker, text = m.group(1), m.group(2)
            style = "List Number" if marker[0].isdigit() else "List Bullet"
            p = doc.add_paragraph(style=style)
            add_inline(p, text)
            i += 1
            continue

        # 引用
        if stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_inline(p, text, size=BODY_SIZE - 1, color=RGBColor(0x40, 0x40, 0x40))
            i += 1
            continue

        # 独立图片 ![alt](path)
        img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", line)
        if img:
            alt, rel = img.group(1), img.group(2)
            img_path = os.path.join(os.path.dirname(os.path.abspath(md_path)), rel)
            if os.path.isfile(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Mm(150))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(cap, f"图：{alt}", size=9, color=RGBColor(0x59, 0x59, 0x59))
            else:
                p = doc.add_paragraph()
                add_inline(p, f"[图片缺失] {alt}（{rel}）", size=9, color=RGBColor(0x99, 0x99, 0x99))
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        set_first_line_chars(p)
        add_inline(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"[OK] 已生成：{docx_path}")


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="Markdown → 中文排版 Word")
    ap.add_argument("input", help="输入 Markdown 文件")
    ap.add_argument("output", nargs="?", help="输出 docx 路径（默认同名 .docx）")
    args = ap.parse_args()
    out = args.output or (os.path.splitext(args.input)[0] + ".docx")
    convert(args.input, out)
